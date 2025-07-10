import os
from flask import Flask, request, render_template, send_file, flash, redirect, url_for
from docx import Document
from werkzeug.utils import secure_filename
from transformers import AutoModelForSeq2SeqLM, AutoTokenizer, pipeline
import torch

# ─── CONFIG ─────────────────────────────────────────────────────────────
UPLOAD_FOLDER = "uploads"
ALLOWED = {"docx", "txt"}
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

app = Flask(__name__)
app.secret_key = os.getenv("FLASK_SECRET_KEY", "change-me")

# Load t5-small at half precision to fit 512MiB RAM
tokenizer = AutoTokenizer.from_pretrained("t5-small")
model = AutoModelForSeq2SeqLM.from_pretrained(
    "t5-small",
    torch_dtype=torch.float16,
    low_cpu_mem_usage=True
)
generator = pipeline(
    "text2text-generation",
    model=model,
    tokenizer=tokenizer,
    device=-1,
    framework="pt"
)

@app.route('/', methods=['GET','POST'])
def index():
    if request.method == 'POST':
        resume_file = request.files.get('resume')
        jd_file     = request.files.get('jd')
        # Validation
        if not resume_file or not resume_file.filename.lower().endswith('.docx'):
            flash('Please upload a .docx resume')
            return redirect(url_for('index'))
        if not jd_file or not jd_file.filename.lower().endswith('.txt'):
            flash('Please upload a .txt job description')
            return redirect(url_for('index'))

        # Save uploads
        rfn = secure_filename(resume_file.filename)
        jfn = secure_filename(jd_file.filename)
        rpath = os.path.join(UPLOAD_FOLDER, rfn)
        jpath = os.path.join(UPLOAD_FOLDER, jfn)
        resume_file.save(rpath)
        jd_file.save(jpath)

        # Extract text
        resume_text = '\n'.join(p.text for p in Document(rpath).paragraphs)
        jd_text     = open(jpath, 'r', encoding='utf-8').read()

        # Prompt for rewriting bullets only
        prompt = (
            "Rewrite only the bullet points in this resume to be ATS-friendly for the job description, "
            "injecting keywords from the JD and using strong action verbs. "
            "Preserve all headings, spacing, and non-bullet content intact.\n\n"
            "Resume:\n" + resume_text + "\n\n"
            "Job Description:\n" + jd_text
        )

        # Generate with limited length
        out = generator(
            prompt,
            max_length=512,
            num_return_sequences=1
        )[0]['generated_text']

        # Rebuild docx preserving non-bullets
        original = Document(rpath)
        new_doc = Document()
        bullets = [ln[2:].strip() for ln in out.splitlines() if ln.strip().startswith('- ')]
        for para in original.paragraphs:
            style = para.style.name
            if style.lower().startswith('list') and bullets:
                new_doc.add_paragraph(bullets.pop(0), style=style)
            else:
                p = new_doc.add_paragraph()
                p.style = style
                p.add_run(para.text)

        out_path = os.path.join(UPLOAD_FOLDER, f"tailored_{rfn}")
        new_doc.save(out_path)

        # Return tailored resume
        return send_file(
            out_path,
            as_attachment=True,
            download_name=f"tailored_{rfn}"
        )

    return render_template('index.html')
